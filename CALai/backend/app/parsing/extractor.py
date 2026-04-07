"""
PDF/DOCX text extraction — from 02_resume_parsing_engine.md.
Primary: PyMuPDF for PDF, python-docx for DOCX.
"""

import io
import fitz  # PyMuPDF
from docx import Document


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF.
    Falls back to empty string if extraction fails.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text("text")
        doc.close()

        # If very little text extracted, it may be a scanned/image PDF
        if len(text.strip()) < 50:
            return text.strip()

        return text
    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {str(e)}")


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to the correct extractor based on file type."""
    file_type = file_type.lower().strip(".")

    if file_type in ("pdf",):
        return extract_pdf_text(file_bytes)
    elif file_type in ("docx", "doc"):
        return extract_docx_text(file_bytes)
    elif file_type in ("txt", "text"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
